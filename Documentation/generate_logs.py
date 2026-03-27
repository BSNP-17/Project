import os
import random
from docx import Document
from docx.shared import Pt
from datetime import date, timedelta

# --- Configuration ---
START_DATE = date(2026, 2, 21)
END_DATE = date(2026, 3, 15)

TEAM = [
    {"name": "B S Niranjan", "role": "Backend", "file": "Log_Niranjan_Feb21_Mar15.docx"},
    {"name": "Subrahmanya I Patgar", "role": "Frontend", "file": "Log_Subrahmanya_Feb21_Mar15.docx"},
    {"name": "Rahul Shanbhag", "role": "Frontend", "file": "Log_Rahul_Feb21_Mar15.docx"},
    {"name": "Vishwas Barker", "role": "Database", "file": "Log_Vishwas_Feb21_Mar15.docx"},
    {"name": "Farhan Faras", "role": "Testing", "file": "Log_Farhan_Feb21_Mar15.docx"},
]

# --- NEW PROJECT-ORIENTED FILLERS ---
# These are added intelligently to ensure the word count reaches ~150 words without sounding repetitive.
TECH_FILLERS = [
    "I made sure to push all my local commits to the remote repository before shutting down for the day.",
    "I checked the application console logs to ensure no silent errors were happening in the background.",
    "I took a few screenshots of today's progress to include in the final project documentation.",
    "I synced up with the rest of the group on WhatsApp to confirm our tasks for tomorrow.",
    "I ran a quick code formatter over the files I touched to maintain our agreed-upon coding standards.",
    "I reviewed the project timeline we submitted in Phase 1 to ensure we aren't falling behind schedule.",
    "I spent some time clearing out old, commented-out code blocks that were making the files hard to read.",
    "I verified the changes across both desktop and mobile views to ensure nothing was broken visually.",
    "I updated the internal task tracker so the testing team knows this specific module is ready for review.",
    "I had to restart the local development server a couple of times because it was caching old API responses."
]

# --- DAILY NARRATIVES ---
DAILY_NARRATIVES = {
    # --- WEEK 1: Feb 21 to Feb 28 ---
    "2026-02-21": {
        "Backend": "I am currently attending the Yuva Aapda Mitra Scheme camp conducted by the NDRF in Karwar, so my coding time is limited today. However, I didn't want the project to stall. During my evening downtime at the camp, I reviewed the pull requests submitted by the frontend team on GitHub. I left some comments on Subrahmanya's code regarding the API endpoint naming conventions.",
        "Frontend": "With Niranjan away at the NDRF camp, Rahul and I focused entirely on the frontend UI polish. I started building the 'Ticket PDF Generation' feature using the `jspdf` and `html2canvas` libraries. It is quite difficult to capture the React component and convert it into a properly formatted A4 PDF. The text kept overlapping when I downloaded the file.",
        "Database": "I worked on creating a robust backup script. Since the frontend team is heavily testing the UI and inserting dummy records, I wanted to make sure our core configuration data is safe. I wrote a MongoDB shell script that exports the `users` and `buses` collections into a secure JSON file every evening. I also cleaned up some dangling ObjectIDs.",
        "Testing": "I started the User Acceptance Testing (UAT) phase. I asked a few friends from another department to try using the TravelEase website on our local network. It was eye-opening. They immediately found the 'Date Format' confusing because it defaulted to US format (MM-DD) instead of Indian format (DD-MM). I logged this critical UX bug for the frontend team."
    },
    "2026-02-23": {
        "Backend": "Still at the NDRF camp in Karwar. After the day's training sessions, I used my phone to research Spring Boot Email Integration. We want the system to send an automated email with the ticket PDF when a booking is confirmed. I read through the `JavaMailSender` documentation and drafted the basic configuration properties we will need to add to `application.properties` once I return.",
        "Frontend": "I finally got the PDF Ticket generation to work correctly. I had to create a hidden HTML `div` that mimics the look of a physical bus ticket, complete with the PNR, travel dates, and a dummy QR code. The `html2canvas` library takes a snapshot of this hidden div and converts it to an image, which is then embedded into the PDF document.",
        "Database": "I analyzed the database performance metrics on MongoDB Atlas. I noticed that the 'Search Bus' query was sometimes taking over 300ms, which is too slow. I realized we were missing an index on the `departureDate` field. I added the index, and the query time dropped to under 20ms. This optimization will be a great talking point for the viva.",
        "Testing": "I verified the PDF Download feature that the frontend team built. I booked a test ticket and downloaded the PDF. The layout looked good, but I noticed the 'Total Fare' field was missing the Rupee symbol (₹). I also checked the PDF on my mobile phone to ensure the file wasn't corrupted. I logged the missing currency symbol issue."
    },
    "2026-02-24": {
        "Backend": "Coordinating from the Karwar camp, I asked Vishwas to temporarily mock the Email Notification API response so the frontend team isn't blocked. I spent an hour reviewing the database indexes Vishwas created yesterday via the Atlas mobile app. I also planned out the final two API endpoints we need: one for resetting passwords and one for fetching user-specific notifications.",
        "Frontend": "I worked on the 'Contact Us' and 'About Us' pages. These are static pages, but they make the application look complete and professional. I added a dummy Google Maps iframe to the Contact page to show the 'TravelEase Headquarters'. I also created a simple contact form, though it currently just shows a success toast message without actually sending an email.",
        "Database": "I supported the frontend team by creating a new collection called `support_tickets`. Even though the Contact form doesn't send emails yet, I set up the schema so we can store user queries in the database. I defined fields for `userName`, `email`, `queryText`, and `status` (Open/Closed). I tested inserting a few records manually.",
        "Testing": "I performed cross-browser testing on the new static pages. I checked the 'About Us' and 'Contact' pages on Google Chrome, Mozilla Firefox, and Microsoft Edge. I found that the Google Maps iframe was overflowing its container on Edge. I reported this CSS flexbox issue. I also verified that all external links on the footer open in a new tab securely."
    },
    "2026-02-25": {
        "Backend": "The NDRF camp activities are intense, but I managed to hop on a quick 15-minute voice call with the team. We discussed the integration strategy for the final phase. I instructed Rahul on how to handle the HTTP interceptors in Axios so that the JWT token is automatically attached to every request, preventing the '401 Unauthorized' errors they were facing.",
        "Frontend": "Following Niranjan's advice, I refactored our API calling logic. Instead of manually attaching the JWT token to every single Axios request, I created an `axiosInterceptor`. This piece of code automatically injects the `Authorization: Bearer <token>` header into all outgoing requests. It cleaned up our codebase significantly and fixed the random logout bugs we were experiencing.",
        "Database": "I started writing the data dictionary for the final report. This is a massive table that details every single field in our database, its data type, and its constraints. For example, documenting that `phoneNumber` is a String of exactly 10 characters, and `totalFare` is a Double. This is a mandatory requirement for the college documentation guidelines.",
        "Testing": "I tested the new Axios Interceptor logic. I intentionally manipulated the JWT token in the browser's local storage to make it invalid. When I tried to access a protected page, the interceptor correctly caught the 401 error, cleared the invalid token, and redirected me to the Login page. This is a huge improvement for our application's security and stability."
    },
    "2026-02-26": {
        "Backend": "The NDRF camp is wrapping up tomorrow. I spent my free time today outlining the 'System Architecture' chapter for our project report. I sketched a rough diagram on paper showing how the React frontend talks to the Spring Boot REST API, which then connects to MongoDB Atlas. I will digitize this using draw.io as soon as I get back home.",
        "Frontend": "I worked on the 'Loading States' today. Previously, when a user clicked 'Search', the screen just froze for a second before the data appeared. I added custom 'Skeleton Loaders' (grey shimmering boxes) that display while the API request is pending. This makes the application feel much faster and gives the user visual feedback that their request is processing.",
        "Database": "I did a comprehensive review of our MongoDB Atlas security settings. I ensured that there are no hardcoded credentials anywhere. I also generated a new set of database passwords just to be safe before we start deploying the application for the final demo. I exported the final schema diagram using MongoDB Compass.",
        "Testing": "I evaluated the new Skeleton Loaders. I used Chrome DevTools to throttle my network speed to 'Slow 3G' to see how the application behaves on a bad connection. The skeleton loaders worked perfectly, preventing the UI from looking broken. I also verified that timeout errors are handled gracefully if the API takes longer than 10 seconds."
    },
    "2026-02-27": {
        "Backend": "Today is the last day of the Yuva Aapda Mitra Scheme camp in Karwar. I spent most of the day traveling back home. I am exhausted, but I booted up my laptop in the evening to pull the latest frontend code from GitHub. I verified that the local development environment is still running perfectly. I am ready to resume full-speed backend development tomorrow.",
        "Frontend": "I focused on cleaning up the React code. I ran `ESLint` across the entire project to find syntax warnings and unused variables. There were over 50 warnings! I spent the day fixing them—mostly removing unused imports and adding missing dependencies to `useEffect` hooks. A clean codebase is essential before we freeze the code for the final evaluation.",
        "Database": "I helped the frontend team with data mocking. They needed to test the UI for a user who has more than 20 past bookings to see if the pagination works. I wrote a quick loop in the MongoDB shell to insert 25 dummy bookings for a specific test user. This allowed them to verify their pagination logic without manual data entry.",
        "Testing": "I performed a full exploratory testing session. I tried to use the application exactly like a real user would, without following a strict script. I found a minor bug: when searching for a bus, if you press the 'Enter' key instead of clicking the 'Search' button, the form reloads the page instead of fetching data. I logged this form submission bug."
    },
    "2026-02-28": {
        "Backend": "I am back at my desk full-time. Today, I tackled the Email Notification system. I added the `spring-boot-starter-mail` dependency. I created a Gmail account specifically for the project (travelease.project@gmail.com) and generated an App Password. I wrote an `EmailService` class that sends a basic 'Booking Confirmed' text email when the transaction is successful. It felt amazing to see the email arrive in my inbox.",
        "Frontend": "I fixed the form submission bug Farhan found yesterday. I added `e.preventDefault()` to the search form's `onSubmit` handler so the page doesn't refresh when the user presses Enter. I also integrated the forgot password UI, adding a screen where users can enter their email to request a reset link. The UI is simple but highly functional.",
        "Database": "I verified the transactional integrity of the new Email feature. I wanted to ensure that if the email fails to send (e.g., due to network issues), the actual booking is still saved in the database, but flagged as 'Email Pending'. I advised Niranjan to run the email sending logic asynchronously so it doesn't block the main database transaction.",
        "Testing": "I tested the new Email Notification feature. I made a booking using my personal email address. The booking was successful, and I received the confirmation email within 5 seconds. I checked the email content to ensure the PNR and travel dates were correct. I noticed the email subject line had a typo, which I reported for a quick fix."
    },

    # --- WEEK 2: Mar 2 to Mar 7 ---
    "2026-03-02": {
        "Backend": "I improved the Email Service by making it send HTML emails instead of plain text. I wrote a simple HTML template inline that includes the TravelEase logo and a neat table showing the passenger details and total fare. I also moved the email sending logic to an `@Async` method so the user doesn't have to wait for the email to send before seeing the success screen.",
        "Frontend": "We started working on the 'Final Polish' for the UI. We audited all the colors and fonts to ensure they strictly follow our brand guidelines. I also worked on the '404 Not Found' page. Instead of a generic browser error, users who navigate to a wrong URL will now see a nice illustration of a lost bus and a button to return to the homepage.",
        "Database": "I prepared the 'Database Deployment' documentation. If the external examiner asks us to set up the project on a new machine from scratch, we need a standard procedure. I wrote a clear guide on how to install MongoDB locally, create the `travelease` database, and import the `init-data.json` file to populate the essential admin credentials and bus routes.",
        "Testing": "I tested the new 404 Error page. I typed random URLs like `/api/fake-page` and `/book/invalid-id`. The application correctly routed me to the custom 404 page every time without crashing. I also re-tested the email feature to confirm that the new HTML format renders correctly on both Gmail and Outlook clients."
    },
    "2026-03-03": {
        "Backend": "I worked on the 'Forgot Password' backend logic. This requires generating a secure, random OTP (One Time Password), saving it to the database with a 10-minute expiration time, and emailing it to the user. I created a new collection/entity called `PasswordResetToken`. I successfully tested generating the token and sending it via the JavaMailSender service.",
        "Frontend": "I connected the 'Forgot Password' UI to Niranjan's new API. When the user enters their email, they are redirected to an 'Enter OTP' screen. I added a countdown timer (10:00 minutes) to the UI to show the token validity. Managing this timer state in React was tricky because `setInterval` behaves weirdly with React hooks, but I resolved it using a custom hook.",
        "Database": "I created the Schema for the `PasswordResetToken` collection. To handle the 10-minute expiration automatically, I utilized MongoDB's TTL (Time-To-Live) index feature. I set an index on the `createdAt` field with `expireAfterSeconds: 600`. This means the database will automatically delete the OTP document after 10 minutes, completely removing the need for backend cleanup tasks.",
        "Testing": "I tested the entire Forgot Password flow. I requested an OTP, received the email, and successfully reset my password. Then, I tried edge cases: I entered a wrong OTP (failed correctly), and I waited 11 minutes to enter the correct OTP (failed correctly due to expiry). The TTL index and backend logic are working perfectly in tandem."
    },
    "2026-03-04": {
        "Backend": "Code Freeze Day! We agreed to stop adding new features today and focus strictly on documentation and fixing any remaining bugs. I did a final review of the Spring Boot code. I removed all unused variables, formatted the code using the IntelliJ auto-formatter, and ensured every public API method had proper Javadoc comments explaining its inputs and outputs.",
        "Frontend": "Frontend Code Freeze! I ran the final production build using `npm run build` to ensure there are no compilation errors. I checked the bundle size and noticed it was a bit large, so I implemented 'Lazy Loading' for the Admin Dashboard routes. This means the admin code is only loaded if the user actually logs in as an admin, making the main site load faster.",
        "Database": "Final Database Audit. I reviewed every collection to ensure there are no spelling mistakes in the field names. I also exported the absolute final version of the database as a JSON dump. This `final_seed.json` file is what we will submit alongside our code on a CD/Pen-drive to the college department.",
        "Testing": "I executed the final full-system regression test suite. I went through all 65 test cases in our Excel sheet one by one. I am happy to report that 64 out of 65 cases passed. The only failure is a minor CSS glitch on extremely small mobile screens, which we documented as a 'Known Issue' in the testing report."
    },
    "2026-03-05": {
        "Backend": "Today was entirely dedicated to writing the Project Report. I took charge of 'Chapter 4: System Implementation'. I pasted screenshots of my backend code structure and explained the MVC architecture. I wrote detailed paragraphs explaining how the JWT Security Filter works and how we implemented the MongoDB transactions. Writing documentation is exhausting but necessary.",
        "Frontend": "I worked heavily on the Project Report today. I wrote 'Chapter 3: System Design'. I created high-quality Data Flow Diagrams (DFD) Level 0, Level 1, and Level 2 using an online drawing tool. I explained how data flows from the React UI components, through the Axios services, to the backend APIs. I also included the UI wireframes we drew in the first month.",
        "Database": "I contributed to the Project Report by writing the 'Database Design' section. I created a clean, professional Entity-Relationship diagram. Even though we used NoSQL, I mapped out the relationships between Users, Bookings, and Buses clearly. I also included tables detailing every field, its data type, and the validation constraints applied in MongoDB Atlas.",
        "Testing": "I drafted 'Chapter 5: Software Testing' for the final report. I explained our testing methodology (Unit, Integration, System testing). I formatted our Excel test cases into neat Word tables. I also wrote the 'Defect Tracking' summary, showing how we identified bugs and resolved them over the course of the project, proving the system's stability."
    },
    "2026-03-06": {
        "Backend": "Report compilation day. We merged all our individual Word documents into one massive Master Document. I spent hours fixing the formatting—ensuring the font is exactly Times New Roman 12pt, with 1.5 line spacing, and that all headings follow the college numbering format (1.1, 1.2.1, etc.). Word formatting is sometimes harder than coding Java.",
        "Frontend": "I worked on formatting the screenshots in the master report. I ensured every image has a proper caption (e.g., 'Figure 4.1: Login Screen') and is centrally aligned. I also wrote the 'Conclusion' and 'Future Enhancements' chapters, mentioning that adding a real Payment Gateway (like Razorpay) and GPS bus tracking are our next planned steps.",
        "Database": "I proofread the entire master document. I looked specifically for grammatical errors and technical inconsistencies. I found a few places where we called the database 'MySQL' by accident from an old template, and corrected them to 'MongoDB'. Consistency in the final report is critical to avoid losing easy marks during the evaluation.",
        "Testing": "I prepared the 'User Manual' appendix for the report. This includes a step-by-step guide with screenshots on how to navigate the application from a user's perspective. I also finalized the 'Bibliography' section, correctly citing all the online tutorials, documentation (Spring IO, React Docs), and textbooks we referenced during development."
    },
    "2026-03-07": {
        "Backend": "We took a printout of the first draft of the report to submit to our project guide for final corrections. While waiting for feedback, I started preparing the PowerPoint presentation for the final Viva. I created slides outlining the backend architecture, emphasizing the use of Spring Security and transactional database updates.",
        "Frontend": "I worked on the UI for the PPT presentation. I chose a clean, professional template that matches our application's blue-and-white theme. I added slides showcasing the 'Before and After' of our UI redesign. I also embedded the screen-recording of our booking flow directly into the presentation so we have a reliable demo fallback.",
        "Database": "I added a dedicated slide in the PPT to explain our Database choices. I prepared speaking notes to justify why MongoDB was the right choice for handling dynamic seat layouts compared to a traditional relational database. I also rehearsed explaining the Aggregation pipeline we used for the admin dashboard.",
        "Testing": "I prepared my speaking parts for the Viva presentation. I will cover the Testing strategies and the challenges we faced. We did a full dry run of the presentation in the library. We timed ourselves and found we were 3 minutes over the allowed limit, so we cut down the introduction to focus more on the technical implementation."
    },

    # --- WEEK 3: Mar 9 to Mar 15 ---
    "2026-03-09": {
        "Backend": "We received feedback from our guide on the report draft. He asked us to add more details about the 'Exception Handling' logic in Chapter 4. I wrote an additional page explaining our `@ControllerAdvice` global exception handler and how it standardizes error responses. I updated the master Word document and regenerated the Table of Contents.",
        "Frontend": "The guide suggested we include a section on 'React Component Lifecycle' in our report to show theoretical knowledge. I wrote a page explaining how we used `useEffect` for data fetching when the component mounts. I also finalized the presentation slides, ensuring the font sizes are large enough to be read clearly on a projector.",
        "Database": "I updated the report based on the guide's feedback. He wanted to see the JSON structure of a 'Booking' document included in the text. I copied a pretty-printed JSON snippet from MongoDB Compass and added it to the Database Design chapter with a detailed explanation of the embedded arrays and references.",
        "Testing": "I proofread the additions made by the rest of the team today to ensure the formatting remained intact. I also checked the Table of Contents and List of Figures to ensure the page numbers synced correctly after adding the new pages. We are preparing to print the final, spiral-bound copies tomorrow."
    },
    "2026-03-10": {
        "Backend": "Today is a big day. We went to the print shop and printed the final 80-page project reports. We got them spiral-bound with the official college cover page. It feels incredibly satisfying to hold the physical manifestation of months of hard work. I also backed up the final project source code to a Google Drive folder.",
        "Frontend": "With the reports printed, we focused entirely on Viva preparation. We sat down and cross-questioned each other. Niranjan asked me how React's Virtual DOM works and why it is faster. I had to practice answering technical questions confidently without hesitating. We are trying to anticipate everything the external examiner might ask.",
        "Database": "Viva preparation. I practiced explaining the difference between NoSQL and SQL databases. I also reviewed the code for the Aggregation pipelines so I can explain exactly what `$match` and `$group` do line-by-line if the examiner asks me to walk through the code. I feel confident in my understanding of the database layer.",
        "Testing": "I prepared for the Viva by reviewing the entire SRS and Testing chapters. If the examiner asks why certain features weren't implemented, my job is to point to the 'Scope of Project' definition in our SRS to defend our choices. We did another mock presentation in an empty classroom to get used to speaking loudly."
    },
    "2026-03-11": {
        "Backend": "We spent the day setting up the 'Production Environment' on the laptop we will use for the final demo. I installed the latest version of Java and ensured the Spring Boot JAR file runs smoothly without any IDE attached. I also configured the Windows Firewall to ensure the local frontend server can hit the backend APIs without CORS issues.",
        "Frontend": "I configured the production build of the React app on the demo laptop. Instead of running `npm start`, I served the optimized `dist` folder using a lightweight HTTP server (`serve -s build`). The application runs blazingly fast in this environment. We did a full run-through of the demo script on this exact machine to ensure no surprises.",
        "Database": "I verified the database connection on the demo laptop. Since we are using MongoDB Atlas (cloud), I had to ensure the college Wi-Fi doesn't block the specific ports needed to connect to the cluster. I also whitelisted the demo laptop's dynamic IP address in the Atlas security dashboard so we don't get locked out during the presentation.",
        "Testing": "I did a final 'Sanity Check' of the application on the demo laptop. I cleared all browser caches, cookies, and local storage. I then executed the primary flow: Register -> Login -> Search Bus -> Book Ticket -> Download PDF. Everything executed flawlessly. The project is officially stable and ready for the examiner's review."
    },
    "2026-03-12": {
        "Backend": "We had a relaxed day today as all the hard work is done. I spent some time reviewing the source code one last time, mostly just admiring the architecture we built. I also helped juniors in the lab who were struggling with their own Java database connectivity issues. Teaching them helped me reinforce my own Viva concepts.",
        "Frontend": "I took a break from coding today. I reviewed the presentation slides one last time and memorized my speaking cues. We decided on the professional dress code for the Viva day to ensure we look like a cohesive team. I feel very proud of the User Interface we managed to deliver for this application.",
        "Database": "No database changes today. I spent the day reading up on advanced MongoDB concepts just in case the examiner tries to test the limits of our knowledge. I learned a bit about database sharding and replication, even though we didn't implement them, it’s good theoretical knowledge to have in my back pocket.",
        "Testing": "I reviewed the printed copies of our Project Report. I highlighted key sections with a sticky note so we can quickly flip to the ER diagram or the Architecture diagram during the Viva if the examiner asks to see them. Being organized will help us look professional and prepared."
    },
    "2026-03-13": {
        "Backend": "Final Mock Viva day. Our project guide conducted a strict mock viva for us in the lab. He grilled me on why I chose JWT over Session Cookies for authentication. I explained the stateless nature of REST APIs and how JWT scales better. He was impressed with the answer but told me to speak a bit louder. Excellent practice session.",
        "Frontend": "During the mock viva, the guide asked me to explain the component hierarchy of our application. I used the whiteboard to draw the React tree, showing how state is passed down as props from the Main Layout to the Bus Cards. He gave us a 9/10 for the UI design, which was a huge confidence booster.",
        "Database": "The guide tested my database knowledge during the mock viva. He asked why we didn't use MySQL for a transactional system like ticket booking. I defended our choice by explaining MongoDB's flexible schema for seat layouts and how we utilized Spring's `@Transactional` to achieve ACID compliance even in NoSQL. He seemed satisfied with the technical justification.",
        "Testing": "I answered questions about our testing methodology during the mock viva. I presented the Excel defect log to the guide to prove we did rigorous manual testing. He advised us to ensure the demo laptop is fully charged and internet-connected before the external examiner arrives next week. We are in the final stretch now."
    },
    "2026-03-14": {
        "Backend": "Rest and review day. I didn't open the laptop today. I just read through the printed project report like a textbook. I mentally traced the flow of data from the `UserController` down to the `UserRepository`. Understanding the full stack deeply is the best preparation for the final exam. I synced with the team to confirm the reporting time for Monday.",
        "Frontend": "I took a mental break today. I reviewed my notes on React Hooks and Context API. I want to be able to explain exactly why we used `useEffect` in the Search component without hesitating. The project is complete, the report is bound, and the presentation is ready. Nothing left to do but wait for the big day.",
        "Database": "I double-checked the MongoDB Atlas dashboard from my phone to ensure the cluster is active and healthy. I also reviewed the schema diagrams in the printed report. I feel a strong sense of accomplishment looking at the complexity of the data structures we successfully implemented and managed over these three months.",
        "Testing": "I packed my bag for the Viva. I kept the pen-drive containing the source code, the presentation, and the digital copy of the report ready. I also kept the spiral-bound reports safely. I sent a final encouraging message to the team WhatsApp group. We built a great project, and we are ready to showcase it."
    }
}

def generate_entry(doc, date_obj, role, member_name):
    date_str = date_obj.strftime("%Y-%m-%d")
    
    # 1. Get Base Narrative
    if date_str in DAILY_NARRATIVES:
        if "Backend" in role: key = "Backend"
        elif "Frontend" in role: key = "Frontend"
        elif "Database" in role: key = "Database"
        else: key = "Testing"
        base_text = DAILY_NARRATIVES[date_str][key]
    else:
        # Fallback for empty days
        base_text = "Today I reviewed the recent modules and focused on final polish and documentation."

    # 2. Add Length to hit ~150 words cleanly
    full_text = base_text
    used_fillers = set()
    safety_count = 0
    
    # Target word count is around 145-160
    while len(full_text.split()) < 145 and safety_count < 5:
        filler = random.choice(TECH_FILLERS)
        if filler not in used_fillers and filler not in full_text:
            full_text += " " + filler
            used_fillers.add(filler)
        safety_count += 1
            
    # 3. Add to Doc
    p_date = doc.add_paragraph()
    runner = p_date.add_run(f"Date: {date_obj.strftime('%d-%m-%Y')}")
    runner.bold = True
    runner.font.size = Pt(12)
    
    p_activity = doc.add_paragraph(full_text)
    p_activity.alignment = 3 # Justify
    doc.add_paragraph("_" * 50)

def generate_student_log(member):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    doc.add_heading('DAILY LOG BOOK (PHASE 4: Final)', 0)
    doc.add_paragraph(f"Name: {member['name']}")
    doc.add_paragraph(f"Role: {member['role']}")
    doc.add_paragraph(f"Project: TravelEase")
    doc.add_page_break()

    current_date = START_DATE
    
    while current_date <= END_DATE:
        if current_date.weekday() != 6: # Skip Sunday
            generate_entry(doc, current_date, member['role'], member['name'])
        current_date += timedelta(days=1)

    doc.save(member['file'])
    print(f"Generated {member['file']}")

if __name__ == "__main__":
    for member in TEAM:
        generate_student_log(member)